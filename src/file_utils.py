import logging
import docx
import threading
from config import Config
import io
from src.s3Bucket import upload_to_s3, Delete_Old_Files_From_S3

logger = logging.getLogger(__name__)

# Check file extension
def allowed_file(filename, allowed_extensions=[]):
    """"Check if a file has an allowed extension"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

# Save Transcription
def save_transcription(transcription, timestamped_filename, srt=None, logger=logger):
    """ Save transcription to a file on S3 """
    logger.info(f"Saving transcription for {timestamped_filename}")
    
    # Save txt file
    txt_content = f"{timestamped_filename}\n{transcription}\n{'-' * 80}\n"
    txt_path = f"{timestamped_filename.rsplit('.', 1)[0]}.txt"
    upload_to_s3(txt_content.encode('utf-8'), txt_path)
    logger.info(f"Transcription saved to S3: {txt_path}")
    
    # Save docx file
    doc = docx.Document()
    doc.add_heading(f'Transcript for {timestamped_filename}', 0)
    doc.add_paragraph(transcription)
    docx_binary = io.BytesIO()
    doc.save(docx_binary)
    docx_binary.seek(0)
    docx_path = f"{timestamped_filename.rsplit('.', 1)[0]}.docx"
    upload_to_s3(docx_binary.getvalue(), docx_path)
    logger.info(f"Transcription saved to S3: {docx_path}")
    
    # Save srt file if it exists
    if srt:
        srt_path = f"{timestamped_filename.rsplit('.', 1)[0]}.srt"
        upload_to_s3(srt.encode('utf-8'), srt_path)
        logger.info(f"SRT file saved to S3: {srt_path}")
    else:
        srt_path = None
    
    return txt_path, docx_path, srt_path

def schedule_cleanup(cleanup_interval=Config.CLEANUP_INTERVAL):
    """Schedule periodic cleanup of old transcript files."""
    Delete_Old_Files_From_S3()
    threading.Timer(cleanup_interval * 60, schedule_cleanup).start()